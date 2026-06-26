from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("docs/whitepaper/L1_Whitepaper.docx")
BODY_FONT = "Latin Modern Roman"
ACCENT_FONT = "Computer Modern Unicode"
FALLBACK_FONT = "Times New Roman"


def set_font(run, size=None, bold=False, italic=False, small_caps=False, spacing=None):
    font = ACCENT_FONT if small_caps else BODY_FONT
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
    run.bold = bold
    run.italic = italic
    run.font.small_caps = small_caps
    if size is not None:
        run.font.size = Pt(size)
    if spacing is not None:
        rpr = run._element.get_or_add_rPr()
        sp = rpr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            rpr.append(sp)
        sp.set(qn("w:val"), str(spacing))


def set_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.42)
    section.right_margin = Inches(1.42)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = True
    sect_pr = section._sectPr
    title_pg = sect_pr.find(qn("w:titlePg"))
    if title_pg is None:
        title_pg = OxmlElement("w:titlePg")
        sect_pr.append(title_pg)


def bottom_rule(paragraph, color="6B6B6B", size="6"):
    ppr = paragraph._p.get_or_add_pPr()
    bdr = ppr.find(qn("w:pBdr"))
    if bdr is None:
        bdr = OxmlElement("w:pBdr")
        ppr.append(bdr)
    bottom = bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run, 9)
    for kind, text in (
        ("begin", None),
        (None, " PAGE "),
        ("separate", None),
        (None, "1"),
        ("end", None),
    ):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        elif text == " PAGE ":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = text
        else:
            node = OxmlElement("w:t")
            node.text = text
        run._r.append(node)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.03
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = styles["Heading 1"]
    h1.font.name = ACCENT_FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), ACCENT_FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), ACCENT_FONT)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = None
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(9)

    h2 = styles["Heading 2"]
    h2.font.name = ACCENT_FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), ACCENT_FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), ACCENT_FONT)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = None
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)

    if "AbstractBody" not in styles:
        s = styles.add_style("AbstractBody", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = BODY_FONT
        s._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
        s.font.size = Pt(10.5)
        s.paragraph_format.left_indent = Inches(0.45)
        s.paragraph_format.right_indent = Inches(0.45)
        s.paragraph_format.line_spacing = 1.02
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "ContentsMajor" not in styles:
        s = styles.add_style("ContentsMajor", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = ACCENT_FONT
        s._element.rPr.rFonts.set(qn("w:ascii"), ACCENT_FONT)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), ACCENT_FONT)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
        s.font.size = Pt(12)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(7)
        s.paragraph_format.space_after = Pt(3)

    if "ContentsMinor" not in styles:
        s = styles.add_style("ContentsMinor", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = BODY_FONT
        s._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
        s.font.size = Pt(10.5)
        s.font.bold = False
        s.paragraph_format.left_indent = Inches(0.27)
        s.paragraph_format.space_after = Pt(1)


def paragraph(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_font(r)
    if align is not None:
        p.alignment = align
    return p


def add_run(p, text, **kwargs):
    r = p.add_run(text)
    set_font(r, **kwargs)
    return r


def set_running_header(section, text):
    section.different_first_page_header_footer = True
    sect_pr = section._sectPr
    title_pg = sect_pr.find(qn("w:titlePg"))
    if title_pg is None:
        title_pg = OxmlElement("w:titlePg")
        sect_pr.append(title_pg)

    first = section.first_page_header.paragraphs[0]
    first.text = ""
    first.paragraph_format.space_after = Pt(0)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(1)
    add_run(hp, text, size=9, small_caps=True, spacing=18)
    bottom_rule(hp)
    add_page_number(section.first_page_footer.paragraphs[0])
    add_page_number(section.footer.paragraphs[0])


def clear_section_links(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False


def title_page(doc):
    for _ in range(3):
        paragraph(doc)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(9)
    add_run(p, "Aetra", size=21, spacing=34)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(14)
    add_run(p, "decentralized proof-of-stake execution network", size=14, spacing=8)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "Daniil Shcherbakov", size=12, spacing=4)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(20)
    add_run(p, "July 8, 2026", size=12, spacing=4)

    for _ in range(2):
        paragraph(doc)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Abstract", size=10.5, bold=True)

    paragraph(
        doc,
        "The aim of this text is to provide a first outline of Aetra: a "
        "decentralized blockchain network designed for trust, moderate speed, "
        "and practical validator operation rather than maximum throughput at "
        "any cost. Aetra targets stronger-than-average decentralization, a "
        "medium hardware profile, bounded validator influence, and native "
        "smart contracts based on the Aetra Virtual Machine.",
        style="AbstractBody",
    )
    paragraph(
        doc,
        "The system is intended to support proof-of-stake validation, nominator "
        "pools, controlled inflation, storage accountability, and future "
        "throughput scaling without making synchronization or validator entry "
        "unnecessarily hard.",
        style="AbstractBody",
    )

    paragraph(doc)
    h = paragraph(doc)
    add_run(h, "Introduction", size=16, bold=True, spacing=12)

    p = paragraph(doc)
    add_run(p, "Aetra", italic=True)
    add_run(
        p,
        " is a blockchain network project built around a conservative "
        "trade-off: it should be faster and more usable than slow settlement "
        "layers, but it should not sacrifice decentralization, auditability, "
        "or validator accessibility in order to compete for the highest "
        "possible transaction count.",
    )

    paragraph(
        doc,
        "The validator economy is an important part of this design. Aetra "
        "should not let the largest holders grow faster only because they "
        "already control more stake. Validator power caps, nominator pools, "
        "commission bounds, concentration metrics, and reward modifiers are "
        "intended to reduce cartel formation and pressure against a small set "
        "of operators.",
    )

    paragraph(
        doc,
        "At the same time, the network must remain economically realistic. "
        "Validators should earn enough to justify reliable infrastructure, "
        "monitoring, and operational risk, but the protocol should avoid "
        "turning high yield into its main product.",
    )

    paragraph(
        doc,
        "Aetra also includes a native virtual machine for smart contracts. "
        "Application logic such as tokens, NFTs, domains, markets, and exchange "
        "contracts should be implemented through AVM standards, while "
        "protocol-critical systems remain native parts of the chain.",
    )


def contents_page(doc):
    doc.add_page_break()
    h = paragraph(doc)
    add_run(h, "Contents", size=16, bold=True, spacing=10)

    items = [
        ("1", "Functional Description of Aetra Components", [
            ("1.1", "Aether Core"),
            ("1.2", "Accounts and Address Policy"),
            ("1.3", "Validator Economy"),
            ("1.4", "Anti-Concentration Policy"),
            ("1.5", "Aetra Virtual Machine"),
            ("1.6", "Storage Rent and State Accountability"),
            ("1.7", "System Entities"),
            ("1.8", "Zones, Routing, and Future Scaling"),
        ]),
        ("2", "System Architecture", [
            ("2.1", "Cosmos SDK and CometBFT Stack"),
            ("2.2", "Three-Tier Module Architecture"),
            ("2.3", "Address System"),
            ("2.4", "Native vs AVM Boundary"),
            ("2.5", "Validator Economy Overview"),
        ]),
        ("3", "Aetra Blockchain", [
            ("3.1", "Accounts, Addresses, and State"),
            ("3.2", "Messages, Transactions, and Receipts"),
            ("3.3", "Genesis, Export, and Import"),
            ("3.4", "Upgrades and Invariants"),
        ]),
        ("4", "Consensus and Staking", [
            ("4.1", "Validator Set and Finality"),
            ("4.2", "Nominator Pools"),
            ("4.3", "Slashing, Evidence, and Insurance"),
            ("4.4", "Anti-Concentration Policy"),
        ]),
        ("5", "Aetra Virtual Machine", [
            ("5.1", "Deploy and Execute Pipeline"),
            ("5.2", "External and Internal Messages"),
            ("5.3", "Gas, Storage Rent, and Exit Codes"),
            ("5.4", "Contract Standards"),
        ]),
        ("6", "Scalability and Network Operation", [
            ("6.1", "Execution Zones"),
            ("6.2", "Scheduling and Load Control"),
            ("6.3", "Experimental Sharding Roadmap"),
            ("6.4", "Validator Hardware Profile"),
        ]),
        ("7", "Native Economy", [
            ("7.1", "AET Supply"),
            ("7.2", "Fees, Burn, Treasury, and Rewards"),
            ("7.3", "Inflation and Validator Income"),
        ]),
        ("8", "Governance and Safety Gates", [
            ("8.1", "Config and Constitution"),
            ("8.2", "Public Testnet Criteria"),
            ("8.3", "Mainnet Criteria"),
        ]),
        ("Conclusion", "Conclusion", []),
        ("A", "The AET Coin", []),
    ]

    for num, title, subs in items:
        p = paragraph(doc, style="ContentsMajor")
        label = title if num == title else f"{num}   {title}"
        add_run(p, label)
        for subnum, subtitle in subs:
            p = paragraph(doc, style="ContentsMinor")
            add_run(p, f"{subnum}   {subtitle}")


def add_chapter_header(section, label):
    clear_section_links(section)
    section.different_first_page_header_footer = False
    hp = section.header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(2)
    add_run(hp, label, size=9.5, small_caps=True, spacing=16)
    bottom_rule(hp)
    fp = section.footer.paragraphs[0]
    fp.text = ""
    add_page_number(fp)


def add_chapter_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.03
    add_run(p, text, size=11)


def chapter_subheading(doc, text):
    p = paragraph(doc)
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, size=11.5, bold=True)


def chapter_one(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page(section)
    add_chapter_header(section, "Chapter 1. Functional Description of Aetra Components")

    h = paragraph(doc)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(10)
    add_run(h, "1    Functional Description of Aetra Components", size=16, spacing=8)

    p = paragraph(doc)
    add_run(p, "Aetra", italic=True)
    add_run(
        p,
        " is a decentralized proof-of-stake execution network implemented on "
        "Cosmos SDK v0.54 with CometBFT v0.39 BFT consensus. The design is "
        "deliberately split into a small native protocol core and a programmable "
        "smart-contract layer called the Aetra Virtual Machine (AVM). That split "
        "is not cosmetic: native keepers protect consensus, validator rules, fee "
        "flow, storage accountability, and upgrade safety, while the AVM exposes "
        "a controlled execution surface for user logic. The network targets "
        "five-to-eight-second block times, finality within five to fifteen "
        "seconds, and a medium hardware validator profile that does not exclude "
        "independent operators. The native asset is AET, denominated in naet, "
        "where one AET equals ten to the ninth naet. This chapter explains how "
        "the major pieces fit together in practice and why each boundary exists.",
    )

    p = paragraph(doc)
    add_run(
        p,
        "The key principle behind the architecture is separation of concerns. "
        "Aetra does not try to turn the base chain into a general-purpose app "
        "runtime. Instead, it keeps the consensus-critical path narrow and "
        "auditable, then layers programmability on top in a way that can be "
        "measured, limited, and disabled without changing the core safety model. "
        "The rest of this chapter describes the native protocol layer, account "
        "rules, validator economics, concentration controls, AVM constraints, "
        "storage rent, reserved system entities, and the scaling roadmap that "
        "is intentionally held behind evidence gates.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Read another way, the chapter is a map of responsibility. The protocol core decides who may participate in consensus, the address layer decides what names are valid, the economic layer decides how stake and fees should behave, and the execution layer decides what user logic may do inside its boundaries. Each piece is narrow on purpose so the chain can stay explainable even as the module count grows.",
    )

    chapter_subheading(doc, "Aether Core")
    p = paragraph(doc)
    add_run(
        p,
        "Aether Core is the native protocol layer that maintains consensus-facing "
        "state, validator lifecycle rules, protocol configuration, fee accounting, "
        "storage accountability, routing commitments, and upgrade safety. It is "
        "implemented as a set of Cosmos SDK keepers with authority-based "
        "governance access control, so the chain can enforce who may mutate a given "
        "module without relying on ad hoc application code. The launch-core surface "
        "includes validator-election, validator-registry, nominator-pool, "
        "single-nominator-pool, native-account, storage-rent, emissions, treasury, "
        "fee-collector, burn, fees, mint-authority, contracts, and "
        "delegator-protection. Together these modules own persistent KV store state "
        "and participate in block processing, but they are still not a place for "
        "ordinary app business logic.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "In functional terms, Aether Core is the chain's control plane. It decides "
        "who can validate, how stake is represented, how fees are accounted for, "
        "where economic value flows, and how system accounts are protected from "
        "user-level mutation. The separation matters because it keeps the protocol "
        "deterministic under load and makes future changes easier to evaluate: if a "
        "new feature cannot be expressed safely in a keeper with explicit authority "
        "rules, it should not be in the core path. That is also why the launch-core "
        "modules are grouped around stable responsibilities rather than around a "
        "single monolithic application object.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The control-plane framing is important because it explains why the native modules are written so conservatively. These keepers are not optimized for novelty; they are optimized for explicit ownership, deterministic state transitions, and clear failure domains. Aether Core is expected to survive network stress, upgrade pressure, and governance change without having to reinterpret what its state means.",
    )

    chapter_subheading(doc, "Accounts and Address Policy")
    p = paragraph(doc)
    add_run(
        p,
        "Aetra defines three address forms in the app/addressing package. "
        "User-facing accounts use the AE Bech32 prefix, are 48 characters long, "
        "and are derived from a 32-byte public key payload. Internal raw addresses "
        "use the 4: prefix for user workchain addressing or the -7: prefix for "
        "protocol system entities. The network reserves more than 30 named system "
        "addresses with identities such as AETConfig, AETConstitution, "
        "AETValidatorRegistry, AETNominatorPool, AETFeeCollector, AETTreasury, "
        "AETStorageRent, AETMintAuthority, AETBurn, AETEmissions, and others. "
        "These are not just labels: they are protocol endpoints with explicit rules "
        "for whether they can hold, receive, or send funds.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Address policy is enforced at transaction admission through "
        "ValidateAnteAddressPolicy. That function inspects message signers, "
        "recipients, and authorities and rejects malformed encodings, zero "
        "addresses, and reserved system addresses used in roles that are not "
        "allowed. The result is a stricter admission boundary than a plain "
        "application would normally have: tooling can rely on the canonical AE "
        "format for users, while the protocol can reserve its internal entities for "
        "governance and block processing. This reduces ambiguity across modules and "
        "prevents user transactions from drifting into system namespace collisions.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The benefit is operational as much as technical. Wallets and explorers can present one clear account format to users, while internal modules can treat system entities as structurally reserved rather than simply \"well-known\". That makes the network harder to misuse accidentally and easier to validate automatically during transaction admission and genesis checks.",
    )

    chapter_subheading(doc, "Validator Economy")
    p = paragraph(doc)
    add_run(
        p,
        "Staking uses nominator pools as the exclusive route. There is no direct "
        "user-to-validator delegation. The nominator-pool module manages pool "
        "creation, share issuance, reward index tracking, unbonding processing, and "
        "stake allocation across validators. The single-nominator-pool module "
        "provides per-validator pool accounting so that validator relationships are "
        "tracked explicitly rather than inferred. In practice this means stake moves "
        "through a bounded accounting layer before it influences validator power, "
        "which gives the protocol more control over reward fairness and exit rules.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Validator funding follows two modes. In solo mode, the operator provides "
        "the full minimum self-stake of one million AET. In pool-backed mode, the "
        "operator provides 400,000 AET self-stake and accepts up to 600,000 AET in "
        "nominator stake. The target validator set size is 128 with a maximum of "
        "300. The validator election window is 100 blocks with a withdrawal cutoff "
        "at 80 blocks, and unbonded stake remains frozen for 1,000 blocks after "
        "exit before unlocking. Those parameters are not decorative: they define the "
        "liquidity and security envelope of the chain and tell operators what kind "
        "of capital commitment is required to participate.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Commission is also bounded rather than free-form. The commission floor is "
        "300 basis points and the ceiling is 2,000 basis points, enforced by the "
        "dynamic-commission module with modifiers from performance and reputation "
        "scores. The intent is to keep validator economics legible while still "
        "leaving room for operators to differentiate service quality. In other "
        "words, Aetra is designed to reward operational reliability, but not at the "
        "cost of making the economics impossible for delegators to reason about.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "This makes validator participation a full lifecycle, not a single staking event. Operators have to register, meet self-stake requirements, stay inside the commission envelope, remain eligible under election rules, and preserve a performance history that does not undermine the reputation score. The chain therefore treats validator power as something that must be maintained rather than simply acquired.",
    )

    chapter_subheading(doc, "Anti-Concentration Policy")
    p = paragraph(doc)
    add_run(
        p,
        "The network does not treat raw stake as the only source of power. The "
        "stake-concentration module measures top-N concentration metrics and caps "
        "effective validator voting power at 300 basis points of total voting power "
        "by default. The validator-score module reduces the reward advantage of "
        "stake that sits beyond the cap, while the aetra-staking-policy module "
        "allocates nominator stake so that no single validator receives more than "
        "300 basis points of total pool allocation and each participant still keeps "
        "a minimum allocation of 25 basis points.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The point of that policy layer is not to punish scale for its own sake. "
        "It is to avoid letting one validator or one stake cluster dominate the "
        "economics of the network simply because it can attract more capital. The "
        "reputation module feeds validator behavior back into fee priority "
        "calculations: a neutral score is 5,000 basis points, low reputation can "
        "incur a premium of up to 500,000 naet, and high reputation can receive a "
        "discount of up to 500,000 naet. That creates a second channel of protocol "
        "feedback, so the chain can reward trustworthy behavior without turning the "
        "base economics into a winner-takes-all market.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The policy matters because concentration risk tends to arrive slowly and then all at once. By keeping the threshold visible and the reward adjustments measurable, Aetra gives governance and operators a chance to react before centralization becomes irreversible. The objective is not absolute equality of stake; it is a validator set that remains resilient, legible, and open to new operators.",
    )

    chapter_subheading(doc, "Aetra Virtual Machine")
    p = paragraph(doc)
    add_run(
        p,
        "The Aetra Virtual Machine is the programmable layer for contracts, but it "
        "is still treated as a bounded subsystem rather than a replacement for the "
        "native protocol. It is implemented as a prototype module gated by genesis "
        "configuration. The AVM supports code storage with a maximum of 4 MB per "
        "contract, contract storage of up to 64 MB per contract, and execution gas "
        "capped at 100 million units per execution. The async layer (aetravm) "
        "limits messages to 32 per transaction and 128 per block, with a maximum "
        "recursion depth of 8 and a maximum of 256 actions per execution.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Those limits are intentional guardrails. Contract deployment costs 1,000 "
        "naet, storage fees are 1 naet per byte, and the AVM handles external and "
        "internal messages, gas accounting, host-function limits, contract "
        "storage, receipts, events, and exit codes. Contracts can be deployed, "
        "executed, upgraded, migrated, and disabled through the contracts keeper "
        "message server interface. The result is a programmable runtime that is "
        "usable enough for real application logic but still small enough to audit "
        "and reason about under block-time pressure.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The hard boundary is that protocol safety remains native. Staking, "
        "slashing, validator election, minting, burning, fee collection, treasury "
        "accounting, storage rent, and system configuration are implemented as "
        "native keepers, not contracts. That keeps the highest-risk logic out of "
        "the contract layer and ensures that the rules controlling supply, "
        "consensus, and network governance are not dependent on third-party code.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "For contract authors, the boundary is a contract with the chain rather than a limitation to work around. They get gas metering, bounded message counts, storage accounting, receipts, and upgrade hooks, but they do not get to rewrite consensus rules or bypass protocol accounting. That split keeps application freedom high while still protecting the parts of the system that must remain stable across the entire network.",
    )

    chapter_subheading(doc, "Storage Rent and State Accountability")
    p = paragraph(doc)
    add_run(
        p,
        "Storage rent is a protocol-level mechanism implemented in the "
        "storage-rent module. The rent rate is one naet per byte per block. The "
        "module defines eleven state classes including wallet, contract, pool "
        "contract, pool share, pool allocation, pool reward index, pool unbonding, "
        "domain record, staking reputation, system module, and validator record. "
        "Contracts with unpaid rent transition through active, frozen, deleted, and "
        "archived states, which gives the chain a way to express state lifecycle "
        "instead of assuming that all data should persist forever at no cost.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The retention period before frozen records are pruned is 10,000 blocks "
        "with a 100-block unfreeze buffer. Collected rent is distributed with 50 "
        "percent to the fee collector, 40 percent to the treasury, and 10 percent "
        "burned. A system rent reserve tracks available funds and projected runway, "
        "generating warnings at critical thresholds. The debt freeze threshold is "
        "set to a minimum runway of 365 days. System module accounts are exempt "
        "from rent freezing so that consensus-critical state remains recoverable "
        "and the protocol cannot accidentally rent-freeze the very accounts it "
        "needs in order to continue operating.",
    )

    chapter_subheading(doc, "System Entities")
    p = paragraph(doc)
    add_run(
        p,
        "Aetra defines over 30 reserved system addresses for protocol "
        "responsibilities that must not behave like ordinary user contracts. These "
        "include the configuration entity (AETConfig), the constitutional bounds "
        "entity (AETConstitution), system registry coordinates (AETSystemRegistry), "
        "validator registry (AETValidatorRegistry), nominator pool "
        "(AETNominatorPool), validator election coordination (AETValidatorElection), "
        "treasury (AETTreasury), fee collector (AETFeeCollector), mint authority "
        "(AETMintAuthority), burn (AETBurn), emissions (AETEmissions), storage rent "
        "(AETStorageRent), validator insurance (AETValidatorInsurance), delegator "
        "protection (AETDelegatorProtection), reputation coordination "
        "(AETReputation), performance tracking (AETPerformance), identity entity "
        "(AETIdentityRoot), bridge hub (AETBridgeHub), sharding coordination "
        "(AETShardingCoordinator), and others.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Each system entity has a defined status indicating whether it is enabled "
        "at genesis, gated by testnet evidence, or reserved for future use. The "
        "important functional point is that these entities cannot be created, "
        "modified, or deleted by user transactions. They represent protocol control "
        "surfaces, not application objects, which is why they are reserved and "
        "tracked explicitly in the codebase rather than exposed as free-form names.",
    )

    chapter_subheading(doc, "Zones, Routing, and Future Scaling")
    p = paragraph(doc)
    add_run(
        p,
        "The x/zones, x/routing, x/load, x/scheduler, and x/sharding modules define "
        "the architecture for future execution zones, deterministic message routing, "
        "load tracking, and shard coordination. These modules are classified as "
        "prototype-only and are not active at genesis. The avm-scheduler module "
        "defines execution batches with dependency graph resolution, with "
        "parameters capping executions per block, parallelism depth, queue depth, "
        "read-set and write-set keys, and receipt count.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The current operating model therefore remains simple: Aetra starts as a "
        "single-shard BFT L1 with AVM support and strong validator economics. Zone "
        "commitments, cross-zone messages, and shard coordination are future-facing "
        "features, and they should only move forward when simulator coverage, "
        "adversarial tests, export/import behavior, long-run testnet evidence, and "
        "independent audit demonstrate that the mechanisms are deterministic and "
        "safe. In this design, scaling is a release gate, not a promise written in "
        "marketing copy.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "That is the reason Chapter 1 keeps returning to boundaries. The chain can only scale safely if the meaning of each component stays stable while the surrounding implementation evolves. If a future zone, routing, or sharding feature cannot be proven deterministic and auditable, it does not belong in the release path yet.",
    )


def chapter_two(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page(section)
    add_chapter_header(section, "Chapter 2. System Architecture")

    h = paragraph(doc)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(10)
    add_run(h, "2    System Architecture", size=16, spacing=8)

    p = paragraph(doc)
    add_run(p, "L1", italic=True)
    add_run(
        p,
        " is built on Cosmos SDK v0.54 with CometBFT v0.39 consensus. "
        "This chapter describes the actual implementation architecture: the "
        "blockchain framework stack, the three-tier module organization, the "
        "custom addressing scheme, the boundary between native protocol code "
        "and smart contracts, and the validator economy as currently implemented.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The implementation is intentionally conventional at the framework layer and intentionally strict at the policy layer. That combination lets the chain inherit mature Cosmos SDK execution primitives while still imposing its own rules for addresses, fees, staking paths, module permissions, and upgrade sequencing.",
    )

    chapter_subheading(doc, "Cosmos SDK and CometBFT Stack")
    p = paragraph(doc)
    add_run(
        p,
        "The network uses Cosmos SDK v0.54 with CometBFT v0.39, providing BFT "
        "finality, ABCI integration, and a modular keeper framework. Target block "
        "time is five to eight seconds with finality within five to fifteen seconds "
        "under normal conditions. The native asset is AET with base denomination "
        "naet where one AET equals ten to the ninth naet. All fee accounting, "
        "emissions, staking rewards, and treasury flows are denominated in naet. "
        "The SDK Bech32 account prefix is ae with BondDenom set to naet.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The fee model uses dynamic deterministic pricing with a default minimum "
        "gas price of zero naet and a minimum gas unit cost of 100,000. "
        "The minimum default fee amount is one naet. The ante handler chain "
        "consists of four layers. The outermost layer is the FeesKeeper ante "
        "handler decorator, which validates fee transactions, checks address "
        "policy, validates transaction envelope limits, and admits or rejects "
        "based on the fee model. The next layer is the "
        "RejectDirectUserStakingDecorator, which prevents users from staking "
        "directly to the SDK staking module. The third layer is the "
        "StorageRentDecorator, which rejects transactions from contracts with "
        "frozen storage rent status. The innermost layer is the SDK base ante "
        "handler with account keeper, bank keeper, sign mode handler, fee grant "
        "keeper, signature verification gas consumer, and unordered transaction "
        "configuration.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The keeper framework follows two patterns: native keepers constructed "
        "with an authority address for governance access control, and persistent "
        "keepers backed by KV stores with full genesis export and import support. "
        "The ABCI lifecycle includes PreBlock for upgrade version map checks, "
        "BeginBlock running SDK modules first then Aetra system modules then "
        "prototype modules, EndBlock processing bank, governance, Aetra modules, "
        "and fee distribution, and InitGenesis for all modules in dependency "
        "order. Fees are split after execution with 98 percent directed to "
        "validator rewards and 2 percent to the community pool.",
    )

    chapter_subheading(doc, "Three-Tier Module Architecture")
    p = paragraph(doc)
    add_run(
        p,
        "Modules are organized into three tiers enforced by a launch module "
        "inventory system embedded in the app package and programmatic wiring "
        "validation. The inventory tracks every module with its x/ directory "
        "path, module name, classification, wiring status, testnet readiness "
        "reason, consensus state ownership, KV-backed runtime mutation flags, "
        "export and import coverage, invariant coverage, and block lifecycle "
        "scanning risk. The application wires 48 modules total: 14 Cosmos SDK "
        "standard modules, 14 Aetra system modules, and 18 Aetra prototype "
        "modules. An additional 14 native economy modules handle fee and supply "
        "operations.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Launch-core modules are consensus-critical and always wired. They own "
        "persistent KV store state and participate in BeginBlock, EndBlock, and "
        "genesis operations. The 14 launch-core modules are: validator-election, "
        "validator-registry, nominator-pool, single-nominator-pool, native-account, "
        "storage-rent, emissions, treasury, fee-collector, burn, fees, "
        "mint-authority, contracts, and delegator-protection. Core modules "
        "maintain dedicated store prefixes and define their own genesis "
        "validation functions. Launch-support modules classified as "
        "launch_support provide governance query surfaces, policy calculations, "
        "and system registries. They are wired into the application but many "
        "are feature-gated and inactive at genesis. Prototype modules classified "
        "as prototype_only or future_avm_standard are gated by the internal "
        "prototype package with a RequireEnabled guard that checks genesis "
        "configuration before allowing message execution. The wiring gate "
        "enforces that prototype modules have no module account permissions, "
        "preventing them from holding or moving native tokens.",
    )

    chapter_subheading(doc, "Address System")
    p = paragraph(doc)
    add_run(
        p,
        "L1 defines a custom addressing scheme in the app/addressing package. "
        "Three address formats serve distinct purposes. User-facing addresses "
        "use the AE prefix with Bech32 encoding, are 48 characters long, and "
        "embed a version byte and magic bytes for format identification. "
        "Internal raw addresses use a 4: prefix (reserved user workchain 4) "
        "with 64 hex characters for bytes-format addressing. Protocol system "
        "addresses use a -7: prefix (reserved system workchain -7) to "
        "distinguish built-in network entities from user accounts. Address "
        "derivation functions in the codec support three roles: account "
        "addresses from user public keys, validator addresses from validator "
        "operator keys, and consensus addresses from consensus public keys.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Malformed addresses are rejected at admission time. Zero-address "
        "abuse is prevented by native-account module checks. The prefix "
        "separation between user workchain 4 and system workchain -7 makes "
        "overlap structurally impossible. Chain-id verification prevents "
        "addresses from different networks from being treated as interchangeable. "
        "Over 30 reserved system addresses are defined with canonical names "
        "such as AETConfig, AETConstitution, AETValidatorRegistry, "
        "AETNominatorPool, AETFeeCollector, AETTreasury, AETStorageRent, "
        "AETMintAuthority, AETBurn, AETEmissions, AETValidatorElection, "
        "AETValidatorInsurance, AETDelegatorProtection, AETReputation, "
        "AETPerformance, AETIdentityRoot, and AETBridgeHub. Each system "
        "address declares flags for fund holding, user fund reception, fund "
        "sending, and genesis activation status.",
    )

    chapter_subheading(doc, "Native vs AVM Boundary")
    p = paragraph(doc)
    add_run(
        p,
        "Protocol safety remains native. Staking, slashing, validator election, "
        "minting, burning, fee collection, treasury accounting, storage rent, "
        "and system configuration are implemented as native Cosmos SDK keepers. "
        "User-defined application logic runs on the AVM, which is gated by "
        "genesis configuration through the internal prototype package. The "
        "wiring gate ensures prototype modules including the contracts module "
        "have no module account permissions, preventing them from holding or "
        "moving native tokens. The project explicitly rejects native asset "
        "module names such as asset and assetfactory. Application-level assets "
        "are contract-only on the AVM. The boundary is enforced at multiple "
        "levels: the ante handler rejects direct staking to the SDK staking "
        "module, the launch module inventory prevents prototype modules from "
        "owning consensus state, and the genesis configuration determines which "
        "prototype features are available at network start.",
    )

    chapter_subheading(doc, "Block Lifecycle and Fee Flow")
    p = paragraph(doc)
    add_run(
        p,
        "The block lifecycle is ordered so that the chain can resolve safety concerns before user transactions are committed. PreBlock runs upgrade version-map checks and other preflight logic, BeginBlock executes the ordered module callbacks for SDK, Aetra system, and prototype modules, EndBlock performs the post-state accounting and validator updates, and InitGenesis defines the starting state in dependency order. That ordering is what makes the application deterministic instead of merely modular.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Fee flow follows the same discipline. The chain keeps fee collection, treasury routing, validator rewards, and burn logic inside native keepers so the economic path is inspectable at every step. This is why the application can talk about a fee model, a treasury, and a reward split in the same breath: all three are part of the same state machine, not separate services bolted onto consensus.",
    )

    chapter_subheading(doc, "Validator Economy Overview")
    p = paragraph(doc)
    add_run(
        p,
        "Staking uses nominator pools as the exclusive route with no direct "
        "user-to-validator delegation. The nominator-pool module manages "
        "up to 10,000 pools, 1,000,000 delegators, and 1,000,000 pending "
        "deposits and withdrawals. The minimum pool deposit is 10 AET. "
        "The unbonding period is 18 days (259,200 blocks at 6-second block "
        "time). Validator funding supports two modes: solo validators needing "
        "one million AET self-stake and pool-backed validators requiring "
        "400,000 AET self-stake with up to 600,000 AET in nominator stake. "
        "Target validator count is 128 with a maximum of 300. The validator "
        "election window is 100 blocks with a withdrawal cutoff at 80 blocks. "
        "Validator voting power is capped at 300 basis points of total voting "
        "power. Commission has a floor of 300 basis points and a ceiling of "
        "2,000 basis points. The aetra-validator-score module applies reward "
        "modifiers based on performance and reputation. Slashing is "
        "evidence-based through the native-evidence module with progressive "
        "penalties for downtime: 1,800 blocks jail for first offense and "
        "14,400 blocks for repeat offenses. Double-sign evidence results in "
        "a slash fraction of 500 basis points with tombstone permanent disable. "
        "Insurance is managed by validator-insurance and delegator protection "
        "by the delegator-protection module with a minimum self-stake ratio "
        "of 4,000 basis points enforced.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The validator layer is also where the chain ties security to accountability. Performance metrics, slashing history, commission discipline, insurance status, and stake concentration all affect how much influence a validator can reasonably carry. The architecture therefore treats consensus power as something that must be continuously justified rather than permanently granted.",
    )


def chapter_three(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page(section)
    add_chapter_header(section, "Chapter 3. Aetra Blockchain")

    h = paragraph(doc)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(10)
    add_run(h, "3    Aetra Blockchain", size=16, spacing=8)

    p = paragraph(doc)
    add_run(p, "This chapter describes the blockchain state model, transaction "
        "lifecycle, genesis initialization, and upgrade mechanisms as implemented "
        "in the current codebase.", italic=False)

    chapter_subheading(doc, "Accounts, Addresses, and State")
    p = paragraph(doc)
    add_run(p, "The native-account module manages all account state on the "
        "network. Each account record includes an AE-prefix Bech32 address, "
        "a 4: prefix or -7: prefix raw address, a balance of AET denominated "
        "in naet, authorization flags, freeze status, storage rent payment "
        "references, and module-specific metadata. Account state is persisted "
        "through the Cosmos SDK KV store abstraction where each module owns a "
        "dedicated store prefix defined by its StoreKey constant. The address "
        "system enforces three address formats: AE-prefix user addresses at 48 "
        "characters, 4: prefix raw user workchain addresses, and -7: prefix "
        "system workchain addresses for over 30 reserved protocol entities. "
        "Address derivation produces three role-dependent variants: account "
        "addresses, validator addresses, and consensus addresses, each derived "
        "from a corresponding public key through the app/addressing Codec.")
    p = paragraph(doc)
    add_run(p, "The native-account module enforces address boundary rules that "
        "prevent user account addresses from colliding with system addresses and "
        "ensure zero addresses are rejected. The launch module inventory tracks "
        "every module with its KV store key, genesis export and import status, "
        "invariant coverage, and block lifecycle participation. Each module's "
        "state schema is validated at genesis time and during export-import "
        "roundtrip testing.")

    chapter_subheading(doc, "Messages, Transactions, and Receipts")
    p = paragraph(doc)
    add_run(p, "Transactions are submitted through CometBFT RPC and processed "
        "through the ABCI CheckTx and DeliverTx path with a four-layer ante "
        "handler decorator chain. The outermost layer is the FeesKeeper ante "
        "handler decorator (x/fees/keeper/ante.go), which validates fee "
        "transactions, checks address policy including zero-address and reserved "
        "system address rejection, validates transaction envelope limits "
        "(maximum 256 KB transaction size, maximum 16 messages per transaction, "
        "maximum 1,024 bytes memo, maximum block gas 20,000,000, maximum 5,000 "
        "transactions per block), and admits or rejects based on the fee model. "
        "The second layer is RejectDirectUserStakingDecorator, which prevents "
        "users from bypassing nominator pools by staking directly to the SDK "
        "staking module. The third layer is StorageRentDecorator, which rejects "
        "transactions originating from contracts with frozen storage rent status. "
        "The innermost layer is the SDK base ante handler performing signature "
        "verification, nonce checking, gas metering through "
        "DefaultSigVerificationGasConsumer, and fee deduction.")
    p = paragraph(doc)
    add_run(p, "The full transfer fee formula includes a minimum transaction "
        "fee, base transfer fee, gas used multiplied by the base fee per gas, "
        "transaction size in bytes multiplied by the byte fee (one naet per byte), "
        "message count multiplied by the message fee (1,000 naet per message), "
        "a congestion surcharge when block utilization exceeds 80 percent, "
        "a low-reputation premium of up to 500,000 naet, a high-reputation "
        "discount of up to 500,000 naet, and storage rent side effects of 100 "
        "naet. Each message type carries the cosmos.msg.v1.signer option "
        "specifying the authorized signer. Messages are routed to the "
        "appropriate keeper by the message router. Failed executions produce "
        "receipts with exit codes, gas used, and error information. Successful "
        "executions commit state changes atomically within the block execution "
        "context.")

    chapter_subheading(doc, "Genesis, Export, and Import")
    p = paragraph(doc)
    add_run(p, "Every network launch starts from a validated genesis state "
        "represented as a map of module names to raw JSON messages. The genesis "
        "state includes initial account balances, validator set membership, "
        "module parameters for all 48 wired modules, over 30 system address "
        "allocations with their fund-holding and fund-sending permissions, "
        "prototype module configuration flags, and default SDK module states "
        "for distribution, governance, mint, protocol pool, and slashing. "
        "Native token metadata for naet and AET is injected into the bank "
        "module genesis. Default parameters are applied through functions in "
        "app/params including AetraMintGenesisState and AetraSlashingParams.")
    p = paragraph(doc)
    add_run(p, "The InitGenesis function for each module is called in "
        "dependency order: Cosmos SDK standard modules first (auth, bank, "
        "staking, distribution, governance, mint, slashing, upgrade, etc.), "
        "followed by Aetra launch-core modules, and then Aetra prototype "
        "modules. InitChain unmarshals the state map, validates through "
        "ValidateGenesis for each module, calls ModuleManager.InitGenesis, "
        "and runs EnsureCoreGenesisCollections to assert system address "
        "integrity. Export and import roundtrip testing validates that the "
        "full chain state can be serialized to JSON and deserialized back "
        "without data loss. The launch module enforces that every persistent "
        "module has export and import coverage. CI validates module account "
        "balances, total supply invariants, validator set consistency, and "
        "system address reservation integrity.")

    chapter_subheading(doc, "Upgrades and Invariants")
    p = paragraph(doc)
    add_run(p, "Protocol upgrades are managed through the Cosmos SDK upgrade "
        "module with cosmovisor support for automated binary switching. Each "
        "upgrade defines a name, target height, and store migration plan. "
        "The current registered upgrade handler is v053-to-v054, which validates "
        "the version map and runs ModuleManager.RunMigrations. A dedicated "
        "native-account version upgrade plan handles account state schema "
        "migration separately. Store upgrades are declared as added module "
        "store keys. The upgrade handler validates that no module versions "
        "regress and that all current modules are present in the version map. "
        "The upgrade playbook documents the process for validators including "
        "proposal submission, voting, binary preparation, and migration "
        "verification.")
    p = paragraph(doc)
    add_run(p, "Invariants are registered at app initialization and run during "
        "block processing to detect state corruption. The invariant set covers "
        "total supply correctness (total supply equals sum of all account "
        "balances plus module accounts), module account balance consistency "
        "(each module account balance matches its recorded state), staking "
        "pool invariant (bonded and unbonded tokens sum correctly), validator "
        "eligibility rules, system address reservation enforcement, and "
        "individual module invariants. Invariant violations halt the chain "
        "with a panic, preventing state divergence across the validator set "
        "and forcing operator intervention before block production continues.")
    p = paragraph(doc)
    add_run(
        p,
        "Those checks are important because the architecture assumes the chain should fail loudly rather than degrade quietly. If a module account drifts, if a validator rule stops matching the registry, or if a storage path becomes inconsistent with the reserve logic, the system should surface that failure immediately. In practice, that makes the whole stack easier to trust because operator response is triggered by a concrete invariant rather than by vague symptoms.",
    )


def chapter_four(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page(section)
    add_chapter_header(section, "Chapter 4. Consensus and Staking")

    h = paragraph(doc)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(10)
    add_run(h, "4    Consensus and Staking", size=16, spacing=8)

    p = paragraph(doc)
    add_run(p, "This chapter describes the consensus mechanism, validator set "
        "management, nominator pool staking, slashing and evidence processing, "
        "and anti-concentration policy as implemented in the current codebase.",
        italic=False)

    chapter_subheading(doc, "Validator Set and Finality")
    p = paragraph(doc)
    add_run(p, "Consensus is provided by CometBFT v0.39, a BFT consensus "
        "engine that provides finality within five to fifteen seconds under "
        "normal conditions. The active validator set is determined by the "
        "validator-election module (x/validator-election), which computes the "
        "set from the validator registry and pool allocations. The election "
        "process runs every 100 blocks with a withdrawal cutoff at 80 blocks. "
        "The maximum active validator set size is 512, with a target of 128 "
        "and a hard maximum of 300 enforced by governance policy. The election "
        "supports up to 10,000 candidates and maintains transition history for "
        "512 epochs. Voting power is computed as min(requested power, self-bond, "
        "max validator power) with additional capping by the stake-concentration "
        "module when the validator exceeds 300 basis points of total voting power.")
    p = paragraph(doc)
    add_run(p, "The validator-registry module (x/validator-registry) maintains "
        "canonical metadata for each validator including operator address, "
        "consensus public key, self-stake amount, pool-backed stake, performance "
        "metrics, jail status, insurance status, and election eligibility. "
        "Validators register with a validator funding mode of solo or pool-backed. "
        "Solo validators require one million AET self-stake. Pool-backed "
        "validators require 400,000 AET self-stake and accept up to 600,000 AET "
        "in nominator stake. Validator funding parameters enforce a minimum "
        "self-stake ratio of 4,000 basis points and a maximum nominator stake "
        "ratio of 6,000 basis points. Frozen stake from exits unlocks after "
        "1,000 blocks.")

    chapter_subheading(doc, "Nominator Pools")
    p = paragraph(doc)
    add_run(p, "There is no direct user-to-validator delegation. All staking "
        "flows through nominator pools managed by the nominator-pool module "
        "(x/nominator-pool) and the single-nominator-pool module "
        "(x/single-nominator-pool). The nominator-pool module supports up to "
        "10,000 pools, 1,000,000 delegators, 1,000,000 pending deposits, "
        "1,000,000 pending withdrawals, and 1,000,000 unbonding entries. "
        "The minimum pool deposit is 10 AET. Pool statuses include active, "
        "paused, frozen limited, and closed. Pool protocol fee is 100 basis "
        "points collected on rewards. The pool protocol fee is split with "
        "5,000 basis points burned, 3,500 basis points to reward fee share, "
        "and 1,500 basis points to treasury fee share.")
    p = paragraph(doc)
    add_run(p, "Users deposit AET into a pool and receive pool shares "
        "representing proportional ownership. Reward indexes accumulate "
        "rewards per share over time. The reward epoch is 14,400 blocks "
        "(one day at 6-second block time). Base reward rate is 350 basis "
        "points with a maximum of 600 basis points. The inflation rate is "
        "bounded between 200 and 600 basis points with a target bonded ratio "
        "of 6,000 basis points. APR targets range from 400 to 700 basis "
        "points. When validators earn rewards through block proposal and "
        "fee collection, rewards are distributed proportionally to pool "
        "participants. Unbonding requests enter an 18-day timelock (259,200 "
        "blocks) during which the stake remains at risk of slashing. "
        "Validator change delay is 100 blocks to allow for orderly "
        "transitions.")

    chapter_subheading(doc, "Slashing, Evidence, and Insurance")
    p = paragraph(doc)
    add_run(p, "Slashing is objective and evidence-based. The native-evidence "
        "module (x/evidence) processes six evidence types: consensus, "
        "double-sign, downtime, missed-block, performance, and fraud. "
        "Evidence submission requires a reporter deposit of 100,000,000 naet "
        "or a fisherman deposit of 25,000,000 naet. Evidence TTL is 10,000 "
        "blocks. The system supports up to 100,000 total evidence records "
        "and 10,000 pending evidence records. Review requires a quorum of "
        "6,700 basis points of voting power. Validator roles include reporter "
        "and fisherman for submitting evidence; only validators vote on "
        "evidence decisions.")
    p = paragraph(doc)
    add_run(p, "Slash fractions are progressive. Downtime first offense "
        "results in a minimum slash fraction of 1 basis point with 1,800 "
        "blocks jail (approximately 3 hours). Repeat offenses multiply the "
        "fraction by 5 and jail for 14,400 blocks (approximately 24 hours). "
        "Chronic offenders (third and subsequent) have the fraction multiplied "
        "by 20. Critical faults including double-sign, consensus, and fraud "
        "use a slash fraction of 500 basis points with tombstone permanent "
        "disable. Frozen stake duration is 259,200 blocks (18 days).")
    p = paragraph(doc)
    add_run(p, "The validator-insurance module (x/validator-insurance) manages "
        "insurance accounting for slashed validators. The delegator-protection "
        "module (x/delegator-protection) enforces minimum self-stake "
        "requirements and limits the impact of validator misbehavior on "
        "nominators. The reporter module (x/reporter) handles reward "
        "accounting for entities that submit valid evidence. Invalid evidence "
        "submissions result in a 5,000 basis point burn of the reporter "
        "deposit with the remaining 5,000 basis points redirected. The "
        "maximum reporter reward is 1,000,000 naet.")

    chapter_subheading(doc, "Anti-Concentration Policy")
    p = paragraph(doc)
    add_run(p, "The stake-concentration module (x/stake-concentration) "
        "measures top-N validator concentration metrics and applies effective "
        "voting-power caps. The default validator power cap is 300 basis "
        "points of total voting power. When a validator exceeds this "
        "threshold, their effective voting power in consensus is capped, "
        "and the aetra-validator-score module (x/aetra-validator-score) "
        "reduces the reward advantage of excess stake beyond the cap. "
        "The aetra-staking-policy module (x/aetra-staking-policy) calculates "
        "pool allocation policy, distributing nominator pool stake across "
        "the validator set with a minimum allocation of 25 basis points "
        "and maximum of 300 basis points per validator.")
    p = paragraph(doc)
    add_run(p, "Validator commission is bounded by governance policy with "
        "a floor of 300 basis points and a ceiling of 2,000 basis points, "
        "enforced by the dynamic-commission module (x/dynamic-commission) "
        "with rate-change limits. The performance module (x/performance) "
        "tracks validator uptime, responsiveness, and historical reliability, "
        "feeding modifiers into commission rates and election scores. The "
        "reputation module (x/reputation) maintains a reputation score for "
        "each validator with a neutral baseline of 5,000 basis points. "
        "Low reputation incurs a fee premium capped at 500,000 naet. High "
        "reputation receives a fee discount capped at 500,000 naet. The "
        "fee priority calculation weighs reputation with 1,000 basis points "
        "and stake priority with 9,000 basis points.")

def chapter_five(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page(section)
    add_chapter_header(section, "Chapter 5. Aetra Virtual Machine")

    h = paragraph(doc)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(10)
    add_run(h, "5    Aetra Virtual Machine", size=16, spacing=8)

    p = paragraph(doc)
    add_run(p, "This chapter describes the Aetra Virtual Machine (AVM), "
        "the smart-contract execution environment, including the deploy and "
        "execute pipeline, internal and external message handling, gas and "
        "storage rent accounting, exit code semantics, the lifecycle matrix, "
        "and contract standard interfaces as implemented in the current codebase.",
        italic=False)

    chapter_subheading(doc, "Deploy and Execute Pipeline")
    p = paragraph(doc)
    add_run(p, "Contract deployment follows a two-step process: StoreCode "
        "followed by DeployContract. StoreCode accepts bytecode prefixed with "
        "the AVM1 magic header, validates it against forbidden non-deterministic "
        "tokens (time.now, wall_clock, random, filesystem, network, float), "
        "computes a canonical code hash using the aetra-avm-code-v1 namespace, "
        "and records the code in the module state. The maximum code size is "
        "4 MB. DeployContract creates a contract instance from stored code with "
        "an init payload of up to 64 KB, a salt of up to 64 KB, and up to 10 "
        "state initialization dependencies. The deployer specifies a creator, "
        "owner, admin, initial balance, upgradeability flag, system ownership "
        "flag, storage schema version, and metadata of up to 1,024 bytes. "
        "The contract address is derived deterministically using the AE Bech32 "
        "user-facing format and the corresponding 4: raw address format. "
        "Each contract record tracks address pair, code ID and hash, state init "
        "hash, creator, owner, admin, upgradeability, system ownership status, "
        "storage schema version, balance, state root, status, storage bytes, "
        "last storage charge height, storage rent debt, logical time, creation "
        "height, and last update height.")
    p = paragraph(doc)
    add_run(p, "External execution is initiated through MsgExecuteExternal, "
        "which specifies the sender, contract address, chain ID, namespace, "
        "optional state init for virtual contracts, payload, funds, gas limit, "
        "metadata, and height. Internal execution is triggered through "
        "MsgExecuteInternal carrying an InternalMessage struct. The contracts "
        "module provides ten message server methods: StoreCode, DeployContract, "
        "ExecuteExternal, ExecuteInternal, SendInternalMessage, "
        "UpgradeContractCode, MigrateContractState, SetContractAdmin, "
        "DisableContractUpgrades, and UpdateContractParams. UpgradeContractCode "
        "requires an actor, contract address, new code ID, and optional "
        "migration handler reference. MigrateContractState transitions between "
        "schema versions. SetContractAdmin changes the admin address. "
        "DisableContractUpgrades permanently removes upgradeability for "
        "contracts that must become immutable.")

    chapter_subheading(doc, "External and Internal Messages")
    p = paragraph(doc)
    add_run(p, "The AVM async execution layer (x/aetravm/async) manages message "
        "queues and execution scheduling with default parameters: 32 messages "
        "per transaction, 128 messages per block, 1,024 queued messages per "
        "contract, 4 maximum processing attempts, 8 maximum recursion depth, "
        "4 KB maximum body size, 64 KB maximum state size, 4 contract deploys "
        "per transaction, 16 deploys per block, 16 emitted messages per "
        "execution, 64 storage writes per execution, 256 actions per execution, "
        "3 retries per message, 1 block default retry delay, 64 blocks maximum "
        "retry delay, and 1,024 dead letter capacity. Execution gas per message "
        "is 10,000 units. Storage fee is 1 naet per byte. Forwarding fee is "
        "1 naet. Contract deployment cost is 1,000 naet.")
    p = paragraph(doc)
    add_run(p, "Internal messages use the InternalMessage struct with fields "
        "for source contract, destination account, funds, opcode, query ID, "
        "body, bounce flag, deadline, gas limit, logical time, message ID, "
        "refunded status, and height. Messages can carry a bounce flag that "
        "triggers a bounce message to the sender on execution failure. "
        "Messages with deadlines are checked for expiry before execution. "
        "The internal message ID is computed deterministically from source, "
        "destination, funds, opcode, query ID, height, bounce, deadline, "
        "logical time, and body. The AVM scheduler (x/avm-scheduler) organizes "
        "execution into batches with dependency graph resolution. Each batch "
        "has a submitted height, initial state root, and task list. The "
        "scheduler groups non-conflicting tasks into parallel execution groups "
        "with a maximum parallelism of 8 tasks per group. Conflicting tasks "
        "with overlapping write sets execute sequentially. Maximum 64 "
        "executions per batch, 128 queue depth, 64 read-set keys and 64 "
        "write-set keys per task, and 1,024 receipts per scheduler state.")

    chapter_subheading(doc, "Gas, Storage Rent, and Exit Codes")
    p = paragraph(doc)
    add_run(p, "Each contract execution has a maximum gas limit of 100,000,000 "
        "units, enforced by the MaxGasPerExecution parameter. The contract "
        "itself has a storage gas cost mechanism: storage rent accrues at "
        "1 naet per byte per block using the StorageRentPerByteBlock parameter. "
        "Contracts track their storage bytes, last storage charge height, "
        "and accumulated storage rent debt. When debt exceeds the available "
        "balance, the contract transitions to frozen status. The lifecycle "
        "matrix defines allowed actions per status. Active contracts can "
        "execute externally, receive internal messages, receive top-ups, "
        "pay rent debt, be queried, emit internal messages, upgrade and "
        "migrate, and be archive-deleted. Frozen contracts accept only "
        "top-ups, rent debt payments, unfreeze actions, and queries. "
        "Frozen-limited contracts add proof queries to the frozen surface. "
        "Archived and deleted contracts are tombstoned and cannot be revived.")
    p = paragraph(doc)
    add_run(p, "Exit codes are organized into four domains. VM execution "
        "errors (codes 0 through 31) include OK, validation failed, "
        "unauthorized, account inactive or frozen, contract frozen, code "
        "rejected, type check failed, out of gas, stack overflow and "
        "underflow, invalid jump, call stack overflow, recursion limit "
        "exceeded, invalid memory access, null reference, corrupted state, "
        "division by zero, arithmetic underflow, gas limit exceeded, and "
        "execution timeout. Action and message errors (32 through 63) "
        "cover expired messages, queue limits, oversized messages, action "
        "budget exceeded, routing failure, queue overflow, and shard "
        "unavailability. State and storage errors (64 through 95) include "
        "storage limit exceeded, storage rent debt, account state too large, "
        "state corruption, state version mismatch, and snapshot failure. "
        "System and host errors (96 through 104) cover execution failure, "
        "internal bounce, forbidden host call, contract abort, assertion "
        "failure, insufficient balance or gas, explicit abort, and inactive "
        "frozen contract.")

    chapter_subheading(doc, "Contract Standards")
    p = paragraph(doc)
    add_run(p, "The contracts module defines a set of standard interfaces "
        "and capabilities that contracts can implement. Asset ownership "
        "records track asset type, contract address, owner address, and "
        "asset ID for token and NFT standards. The native staking "
        "injection capability allows governance-authorized contracts to "
        "interact with nominator pools by sending staking-injection messages "
        "that reference a pool ID and amount. Contracts can receive the "
        "native_staking_injection capability through a governance grant, "
        "which records the contract address pair, pool ID, and grant height. "
        "Contract events are emitted during lifecycle operations with event "
        "types for code storage, contract instantiation, and contract "
        "execution. The codec interface (MsgServer and QueryServer) provides "
        "standard gRPC query methods for listing codes, contracts, storage "
        "entries, receipts, internal message queues, events, and state roots. "
        "The contracts module participates in the system-wide state root "
        "commitment scheme through the RootContribution interface, producing "
        "a deterministic commitment hash that anchors the entire contract "
        "state tree into the protocol state root.")


def chapter_six(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page(section)
    add_chapter_header(section, "Chapter 6. Scalability and Network Operation")

    h = paragraph(doc)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(10)
    add_run(h, "6    Scalability and Network Operation", size=16, spacing=8)

    p = paragraph(doc)
    add_run(p, "This chapter describes the mechanisms for scaling execution "
        "through zones and sharding, scheduling and load control, the "
        "experimental sharding roadmap, and the validator hardware profile "
        "requirements as implemented or planned in the current codebase.",
        italic=False)

    chapter_subheading(doc, "Execution Zones")
    p = paragraph(doc)
    add_run(p, "The zones module (x/zones) defines a zone-based execution "
        "model where each zone represents an isolated execution domain with "
        "its own state prefix, VM policy, fee policy, upgrade policy, data "
        "availability policy, and audit status. Four zone kinds are defined: "
        "FINANCIAL (FINANCIAL_ZONE), IDENTITY (IDENTITY_ZONE), APPLICATION "
        "(APPLICATION_ZONE), and CONTRACT (CONTRACT_ZONE). Each zone declares "
        "a VM policy of AVM, COSMWASM_GATED, or NATIVE_MODULE, an upgrade "
        "policy of GOVERNANCE, SCHEDULED, or IMMUTABLE, and a data "
        "availability policy of CORE_COMMITMENT or REPLICATED. Zone audit "
        "status progresses through EXPERIMENTAL, INTERNAL_REVIEW, and AUDITED "
        "states. Zones are registered by governance with a specified "
        "activation height and become active once that height is reached.")
    p = paragraph(doc)
    add_run(p, "The execution zone model specification defines eight surfaces "
        "that each zone must satisfy. The keeper-set surface requires that "
        "zone adapters avoid direct cross-zone mutations. The state-prefix "
        "surface requires prefix-isolated exportable state. The mempool-policy "
        "surface requires local admission with global bounds. The shard-layout "
        "surface requires committed layout with epoch routing. The message "
        "queues surface requires committed inbox-outbox effects. The fee "
        "policy surface requires zone-local accounting with global settlement. "
        "The proof root surface requires state, message, receipt, and event "
        "domain roots. The execution metrics surface requires zone execution "
        "summary inputs. Each zone model spec is hashed deterministically "
        "using ComputeExecutionZoneModelHash to produce a boundary hash that "
        "anchors the zone specification. Zone commitments are appended to "
        "the zone registry state with monotonically increasing heights and "
        "previous commitment hash chaining.")

    chapter_subheading(doc, "Scheduling and Load Control")
    p = paragraph(doc)
    add_run(p, "The scheduler module (x/scheduler) manages periodic, delayed, "
        "and epoch-based job execution with a maximum of 25 jobs per block, "
        "2,000,000 total scheduler gas per block, and 100,000 maximum gas per "
        "job. Authorized modules include aetracore, fees, load, mesh, "
        "networking, payments, routing, zones, epoch, avm-dex-contract, and "
        "contract-assets. Jobs support retry policies with configurable max "
        "retries and backoff intervals. History retention is 200 records. "
        "Planning supports four modes: sequential execution of all tasks, "
        "optimistic parallel placement of non-conflicting tasks into the same "
        "batch, DAG-based dependency resolution, and mailbox-based scheduling "
        "where tasks are grouped by actor. Conflict detection compares write "
        "sets against read and write sets of other tasks to determine "
        "parallelism safety. When conflicts prevent optimistic parallel "
        "placement, the scheduler falls back to sequential execution with a "
        "conflict_fallback_sequential status.")
    p = paragraph(doc)
    add_run(p, "The load module (x/load) implements EMA-based load scoring "
        "with a 60-block window and exponential smoothing. The default alpha "
        "is 2 over 61 with a maximum delta of 500 basis points per update. "
        "Load is classified into three bands: LOW (below 3,000 basis points), "
        "MEDIUM (3,000 to 7,000 basis points), and HIGH (above 7,000 basis "
        "points). Targets include mempool size of 10,000 transactions, block "
        "gas of 20,000,000, latency of 5 blocks, and execution steps of "
        "20,000,000. Five weighted metrics contribute to the composite score: "
        "mempool size (2,000 basis points), block utilization (3,000 basis "
        "points), transaction latency (2,000 basis points), failure rate "
        "(1,000 basis points), and execution time (2,000 basis points). "
        "The AVM scheduler (x/avm-scheduler) provides task-level scheduling "
        "with batch dependency graphs, parallel execution groups, conflict "
        "counters per contract address, and serial state root computation "
        "for finalized batches.")

    chapter_subheading(doc, "Experimental Sharding Roadmap")
    p = paragraph(doc)
    add_run(p, "The sharding coordinator module (x/sharding-coordinator) "
        "defines parameters for future multi-shard operation. Maximum shards "
        "is 4,096 with 512 validators per shard and 2 minimum validator "
        "coverage. Each validator can serve up to 4 shard assignments. "
        "Shard statuses include paused, active, draining, and disabled. Three "
        "security levels are defined: standard, high, and critical. "
        "Cross-shard routing supports configurable per-route parameters "
        "including enable flags, message byte limits, timeout blocks, and "
        "in-flight limits with a maximum of 1,000,000 timeout blocks and "
        "100,000 in-flight messages. Rebalance proposals allow validators "
        "to be moved between shards with sequence-ordered moves when load "
        "imbalance is detected. Each shard produces state root references "
        "at regular heights that are anchored into the coordinator state. "
        "Load metrics per shard track transactions per block, gas per block, "
        "state bytes, and pending messages with bounds of 10,000,000 "
        "transactions, 1,000,000,000 gas, and 10,000,000 pending messages.")
    p = paragraph(doc)
    add_run(p, "The sharding simulator (x/sharding/sim) provides a framework "
        "for testing shard state transitions, cross-shard message passing, "
        "and masterchain state coordination without requiring a live network. "
        "The roadmap requires that sharding mechanisms are not activated at "
        "genesis. They will be enabled only when simulator coverage, "
        "adversarial tests, export-import roundtrip behavior, long-run "
        "testnet evidence, and independent audit demonstrate that "
        "deterministic safety and liveness guarantees hold. The initial "
        "network operates as a single-shard L1 with AVM support. Multi-shard "
        "operation, zone commitments, cross-zone message routing, and shard "
        "coordination are added incrementally as the evidence base grows.")

    chapter_subheading(doc, "Validator Hardware Profile")
    p = paragraph(doc)
    add_run(p, "The network targets a medium-hardware validator profile. "
        "Block times of five to eight seconds with finality within five to "
        "fifteen seconds require consumer-grade server hardware rather than "
        "enterprise-class machines. State size is bounded through storage "
        "rent with a target of keeping full-node synchronization practical "
        "for independent operators. The 64 MB per-contract storage limit, "
        "storage rent rate of 1 naet per byte per block, and 10,000-block "
        "retention period for frozen contract data prevent unbounded state "
        "growth. The maximum block gas of 20,000,000 limits computational "
        "load per block. The maximum transaction size of 256 KB and 5,000 "
        "transactions per block cap network bandwidth requirements. The "
        "maximum validator set of 300 with a target of 128 balances consensus "
        "communication overhead against decentralization. The validator "
        "change delay of 100 blocks and frozen stake unlock of 1,000 blocks "
        "provide orderly transitions without requiring rapid state "
        "synchronization. These parameters are engineering constraints "
        "designed to keep node operation accessible to technically competent "
        "individuals rather than exclusively to large infrastructure "
        "providers.")


def main():
    doc = Document()
    set_page(doc.sections[0])
    configure_styles(doc)
    set_running_header(doc.sections[0], "Introduction")
    title_page(doc)
    contents_page(doc)
    chapter_one(doc)
    chapter_two(doc)
    chapter_three(doc)
    chapter_four(doc)
    chapter_five(doc)
    chapter_six(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
